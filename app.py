import streamlit as st
import google.generativeai as genai
import midtransclient
import uuid
import os
import sqlite3
import datetime
import pandas as pd
from dotenv import load_dotenv
import docx
from io import BytesIO

# ==========================================
# 1. LOAD CONFIGURATION & SECURITY
# ==========================================
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
MIDTRANS_SERVER_KEY = os.getenv("MIDTRANS_SERVER_KEY")

st.set_page_config(page_title="RPP Merdeka Generator", page_icon="📚", layout="wide")

# ==========================================
# 2. KONFIGURASI DATABASE SQLITE
# ==========================================
def init_db():
    conn = sqlite3.connect('rpp_logs.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS log_transaksi
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  waktu TEXT, order_id TEXT, mata_pelajaran TEXT,
                  materi TEXT, tipe_rpp TEXT, status TEXT)''')
    conn.commit()
    conn.close()

def simpan_log(order_id, mapel, materi, tipe_rpp, status):
    conn = sqlite3.connect('rpp_logs.db')
    c = conn.cursor()
    waktu_sekarang = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("INSERT INTO log_transaksi (waktu, order_id, mata_pelajaran, materi, tipe_rpp, status) VALUES (?, ?, ?, ?, ?, ?)",
              (waktu_sekarang, order_id, mapel, materi, tipe_rpp, status))
    conn.commit()
    conn.close()

def ambil_semua_log():
    conn = sqlite3.connect('rpp_logs.db')
    df = pd.read_sql_query("SELECT * FROM log_transaksi ORDER BY id DESC", conn)
    conn.close()
    return df

init_db()

# ==========================================
# 3. MANAJEMEN SESSION STATE
# ==========================================
if 'order_id' not in st.session_state: st.session_state.order_id = None
if 'payment_url' not in st.session_state: st.session_state.payment_url = None
if 'data_rpp' not in st.session_state: st.session_state.data_rpp = {}
if 'hasil_rpp' not in st.session_state: st.session_state.hasil_rpp = None
if 'log_tersimpan' not in st.session_state: st.session_state.log_tersimpan = False
if 'harga_tagihan' not in st.session_state: st.session_state.harga_tagihan = 0

# ==========================================
# 4. FUNGSI BISNIS
# ==========================================
def buat_link_pembayaran(harga, order_id):
    snap = midtransclient.Snap(is_production=False, server_key=MIDTRANS_SERVER_KEY)
    param = {
        "transaction_details": {"order_id": order_id, "gross_amount": harga},
        "customer_details": {"first_name": "Guru", "email": "guru@contoh.com"}
    }
    transaction = snap.create_transaction(param)
    return transaction['redirect_url']

def cek_status_midtrans(order_id):
    try:
        snap = midtransclient.Snap(is_production=False, server_key=MIDTRANS_SERVER_KEY)
        status_res = snap.transactions.status(order_id)
        return status_res.get('transaction_status')
    except Exception as e: return None

def generate_rpp_ai(data):
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    if data['tipe'] == 'Standar':
        prompt = f"""Anda ahli pendidikan. Buat RPP Kurikulum Merdeka (Standar) untuk:
        - Mapel: {data['mapel']}, Fase: {data['fase']}, Materi: {data['materi']}, Waktu: {data['waktu']}, Tujuan: {data['tujuan']}.
        Struktur: Informasi Umum, Pendahuluan, Inti, Penutup, Asesmen Dasar. Gunakan Markdown."""
    else:
        prompt = f"""Bertindak sebagai Ahli Kurikulum. Buatkan Modul Ajar Kurikulum Merdeka SANGAT DETAIL:
        - Mapel: {data['mapel']}, Fase: {data['fase']}, Materi: {data['materi']}, Waktu: {data['waktu']}, Tujuan: {data['tujuan']}.
        WAJIB: 1. Informasi Umum, 2. Komponen Inti (Tujuan, Pemahaman Bermakna, Pertanyaan Pemantik), 3. Kegiatan Pembelajaran (Pendahuluan, Inti berdiferensiasi, Penutup), 4. Asesmen (3 Soal HOTS + Kunci, Rubrik), 5. Lampiran (Draf LKPD, Pengayaan/Remedial). Gunakan Markdown."""
        
    response = model.generate_content(prompt)
    return response.text

def export_ke_docx(teks):
    doc = docx.Document()
    doc.add_heading('Modul Ajar / RPP Kurikulum Merdeka', level=1)
    for baris in teks.split('\n'):
        baris_bersih = baris.strip()
        if not baris_bersih: continue
        if baris_bersih.startswith('###'): doc.add_heading(baris_bersih.replace('###', '').strip(), level=3)
        elif baris_bersih.startswith('##'): doc.add_heading(baris_bersih.replace('##', '').strip(), level=2)
        elif baris_bersih.startswith('#'): doc.add_heading(baris_bersih.replace('#', '').strip(), level=1)
        else:
            baris_bersih = baris_bersih.replace('**', '')
            doc.add_paragraph(baris_bersih)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# ==========================================
# 5. UI & ALUR APLIKASI
# ==========================================
with st.sidebar:
    st.header("📊 Admin Dashboard")
    if st.button("Tampilkan Log Transaksi"): st.write(ambil_semua_log())

st.title("📚 RPP Kurikulum Merdeka Generator")
st.write("Hasilkan Modul Ajar berkualitas menggunakan AI.")

if st.session_state.order_id is None:
    with st.form("rpp_form"):
        mata_pelajaran = st.text_input("Mata Pelajaran:")
        fase_kelas = st.text_input("Fase / Kelas:")
        materi_pokok = st.text_input("Materi Pokok:")
        alokasi_waktu = st.text_input("Alokasi Waktu:")
        tujuan_pembelajaran = st.text_area("Tujuan Pembelajaran:")
        pilihan_tipe = st.radio("Pilih layanan:", ["RPP Standar (Rp 5.000)", "Modul Ajar PRO (Rp 10.000)"])
        submit_button = st.form_submit_button("Lanjut ke Pembayaran")

    if submit_button:
        if not all([mata_pelajaran, fase_kelas, materi_pokok, alokasi_waktu, tujuan_pembelajaran]):
            st.warning("Lengkapi semua data!")
        else:
            harga_final = 5000 if "Standar" in pilihan_tipe else 10000
            st.session_state.order_id = f"RPP-{str(uuid.uuid4().hex[:8]).upper()}"
            st.session_state.harga_tagihan = harga_final
            st.session_state.data_rpp = {
                "mapel": mata_pelajaran, "fase": fase_kelas, "materi": materi_pokok, 
                "waktu": alokasi_waktu, "tujuan": tujuan_pembelajaran,
                "tipe": "Standar" if "Standar" in pilihan_tipe else "Pro"
            }
            with st.spinner('Menghubungkan ke pembayaran...'):
                st.session_state.payment_url = buat_link_pembayaran(harga_final, st.session_state.order_id)
            st.rerun()

else:
    if st.session_state.hasil_rpp is not None:
        st.success("Pembayaran Berhasil!")
        st.subheader(f"📚 HASIL {st.session_state.data_rpp['tipe'].upper()} ANDA")
        st.markdown(st.session_state.hasil_rpp)
        file_docx = export_ke_docx(st.session_state.hasil_rpp)
        st.download_button("📥 Unduh RPP (.docx)", file_docx, f"RPP_{st.session_state.data_rpp['materi']}.docx")
        if st.button("🔄 Buat Baru"):
            st.session_state.update({'order_id': None, 'hasil_rpp': None, 'log_tersimpan': False})
            st.rerun()
    else:
        st.info(f"Tagihan ID: **{st.session_state.order_id}** | Total: **Rp {st.session_state.harga_tagihan}**")
        st.markdown(f'<a href="{st.session_state.payment_url}" target="_blank"><button style="background-color:#4CAF50; color:white; padding:12px 30px; border:none; border-radius:4px; cursor:pointer;">👉 KLIK BAYAR</button></a>', unsafe_allow_html=True)
        if st.button("🔄 Cek Pembayaran"):
            status = cek_status_midtrans(st.session_state.order_id)
            if status in ['settlement', 'capture']:
                with st.spinner("AI sedang menyusun RPP..."):
                    st.session_state.hasil_rpp = generate_rpp_ai(st.session_state.data_rpp)
                    if not st.session_state.log_tersimpan:
                        simpan_log(st.session_state.order_id, st.session_state.data_rpp['mapel'], st.session_state.data_rpp['materi'], st.session_state.data_rpp['tipe'], "Berhasil")
                        st.session_state.log_tersimpan = True
                    st.rerun()
            else: st.warning("Pembayaran belum terdeteksi.")
